import whisper
import noisereduce as nr
import numpy as np
from pydub import AudioSegment
from docx import Document
from transformers import pipeline
import tkinter as tk
from tkinter import scrolledtext
from tkinter import messagebox
from tkinter import filedialog, scrolledtext, messagebox, END


# Load Whisper model
model = whisper.load_model("base")  # You can choose a larger model for better accuracy

# Function to reduce noise in audio
def reduce_noise_in_audio(audio):
    audio_data = np.array(audio.get_array_of_samples())
    reduced_noise = nr.reduce_noise(y=audio_data, sr=audio.frame_rate)
    reduced_audio = audio._spawn(reduced_noise.tobytes())
    return reduced_audio

# Function to split audio into fixed intervals
def split_audio_fixed_intervals(file_path, chunk_length_ms=60000):
    try:
        audio = AudioSegment.from_file(file_path)
        audio = reduce_noise_in_audio(audio)  # Apply noise reduction
        chunks = []

        # Split into fixed intervals (e.g., 1-minute chunks)
        for i in range(0, len(audio), chunk_length_ms):
            chunk = audio[i:i + chunk_length_ms]
            chunk_file = f"chunk_{i // chunk_length_ms}.wav"
            chunk.export(chunk_file, format="wav")
            chunks.append(chunk_file)

        return chunks
    except Exception as e:
        print(f"Error splitting audio: {e}")
        return []

# Function to transcribe a single audio chunk using Whisper
def transcribe_chunk_with_whisper(chunk_file):
    transcription = model.transcribe(chunk_file)
    return transcription['text']

# Function to transcribe audio and save to a Word document
def transcribe_audio_to_word(file_path, word_doc_path, chunk_length_ms=60000):
    print("Splitting audio into chunks...")
    chunks = split_audio_fixed_intervals(file_path, chunk_length_ms)

    if not chunks:
        return "Error: No chunks were created."
    
    doc = Document()
    doc.add_heading("Audio Transcription", level=1)

    full_transcription = []
    for i, chunk in enumerate(chunks, start=1):
        print(f"Transcribing chunk {i}/{len(chunks)}...")
        transcription = transcribe_chunk_with_whisper(chunk)
        print(f"Chunk {i} Transcription: {transcription}")

        doc.add_paragraph(f"{transcription}")
        full_transcription.append(transcription)

    doc.save(word_doc_path)
    print(f"Transcriptions saved to: {word_doc_path}")
    return word_doc_path

# Function to read text from a Word document
def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Summarize the text using the Huggingface summarizer pipeline
summarizer = pipeline("summarization", model="facebook/bart-large-cnn", tokenizer="facebook/bart-large-cnn", framework="pt")

def summarize_text(text):
    summary = summarizer(text, max_length=300, min_length=50, do_sample=False)
    return summary[0]['summary_text']

# Function to summarize text and format it as bullet points
def summarize_text_as_bullets(text, chunk_size=1024):  # Adjust chunk_size based on model's token limit
    # Split the text into chunks that fit within the model's token limit
    text_chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
    
    bullet_points = []
    
    # Summarize each chunk individually and convert it into bullet points
    for chunk in text_chunks:
        summary = summarizer(chunk, max_length=300, min_length=50, do_sample=False)
        summarized_text = summary[0]['summary_text']
        
        # Split the summarized text into sentences or key points
        points = summarized_text.split('. ')  # Split at each sentence (you can adjust this based on summary structure)
        
        # Add bullet points to each line
        bullet_points.extend([f"• {point.strip()}" for point in points if point.strip()])
    
    # Join all the bullet points together
    return "\n".join(bullet_points)

# Function to process transcription and summarization
def process_transcription_and_summarization(audio_file, word_output, chunk_length_ms=60000):
    transcribe_audio_to_word(audio_file, word_output, chunk_length_ms)

    # Read the transcribed text from the Word document
    transcription_text = read_docx(word_output)

    # Summarize the transcription
    summarized_text = summarize_text_as_bullets(transcription_text)

    return summarized_text

# Tkinter GUI setup
def create_gui():
    def choose_file():
        """Open a file dialog for selecting a .wav file."""
        file_path = filedialog.askopenfilename(
            title="Choose a WAV File",
            filetypes=[("WAV Files", "*.wav")],
        )
        if file_path:
            audio_file_entry.delete(0, END)  # Clear the current text in the entry field
            audio_file_entry.insert(0, file_path)  # Insert the selected file path

    def on_process():
        """Process the audio file and generate the summarized transcription."""
        audio_file = audio_file_entry.get()  # Get the audio file path
        word_output = word_output_entry.get()  # Get the Word document output path

        if not audio_file or not word_output:
            messagebox.showerror("Error", "Both file path and Word document name are required!")
            return

        try:
            result_text.delete(1.0, tk.END)  # Clear the result box
            result_text.insert(tk.END, "Processing, please wait....!!")
            result_text.update_idletasks()
            summarized_text = process_transcription_and_summarization(audio_file, word_output)
            result_text.delete(1.0, END)
            result_text.insert(END, summarized_text)
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

    def on_close():
        """Handle window close event."""
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            window.destroy()

    # Initialize Tkinter window
    window = tk.Tk()
    window.title("Minutes Of Meeting Generator")
    window.geometry("600x400")  # Adjust window size
    window.protocol("WM_DELETE_WINDOW", on_close)  # Handle the close button

    # Audio file input with "Choose File" button
    tk.Label(window, text="Audio File Path:").pack(pady=5)
    file_frame = tk.Frame(window)
    file_frame.pack(pady=5)
    audio_file_entry = tk.Entry(file_frame, width=50)
    audio_file_entry.pack(side=tk.LEFT, padx=5)
    choose_file_button = tk.Button(file_frame, text="Choose File", command=choose_file)
    choose_file_button.pack(side=tk.LEFT)

    # Word output file input
    tk.Label(window, text="Word Document Name:").pack(pady=5)
    word_output_entry = tk.Entry(window, width=50)
    word_output_entry.pack(pady=5)

    # Process Button
    process_button = tk.Button(window, text="Process and Summarize", command=on_process)
    process_button.pack(pady=10)

    # Result text box
    result_text = scrolledtext.ScrolledText(window, width=70, height=20)
    result_text.pack(pady=10)

    window.mainloop()


# Run the GUI
create_gui()
